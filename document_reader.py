import io
import pypdf
from docx import Document
import openpyxl


def clean(text):
    if not text:
        return ""
    return text.encode('ascii', 'ignore').decode('ascii').strip()


def read_pdf(file_bytes: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    text = ""
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text += extracted + "\n"
    return clean(text)


def read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    text = ""
    for para in doc.paragraphs:
        if para.text.strip():
            text += para.text + "\n"
    # Also read tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text += row_text + "\n"
    return clean(text)


def read_excel(file_bytes: bytes) -> str:
    wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True)
    text = ""
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text += f"\n[Sheet: {sheet_name}]\n"
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) for cell in row if cell is not None]
            if row_values:
                text += " | ".join(row_values) + "\n"
    return clean(text)


def extract_text(uploaded_file) -> tuple[str, str]:
    """
    Takes a Streamlit uploaded file object.
    Returns (extracted_text, file_type).
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return read_pdf(file_bytes), "PDF"
    elif name.endswith(".docx"):
        return read_docx(file_bytes), "Word"
    elif name.endswith(".xlsx") or name.endswith(".xls"):
        return read_excel(file_bytes), "Excel"
    else:
        raise ValueError(f"Unsupported file type: {uploaded_file.name}")
