from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def clean(text):
    if not text:
        return ""
    return str(text).encode('ascii', 'ignore').decode('ascii')


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_cell_text(cell, text, bold=False, font_size=10, color=None):
    para = cell.paragraphs[0]
    run = para.add_run(clean(str(text or "")))
    run.bold = bold
    run.font.size = Pt(font_size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(clean(text).upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string("1F3864")


def add_body(doc, text, italic=False, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(clean(str(text or "")))
    run.font.size = Pt(10)
    run.font.name = "Arial"
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def add_bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    run = p.runs[0] if p.runs else p.add_run()
    run.text = clean(str(text or ""))
    run.font.size = Pt(10)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_two_col_table(doc, rows, col_widths=(2.0, 4.5)):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(col_widths[0])
    table.columns[1].width = Inches(col_widths[1])
    for label, value, shade in rows:
        row = table.add_row()
        if shade:
            set_cell_bg(row.cells[0], shade)
        add_cell_text(row.cells[0], label, bold=True, font_size=10)
        add_cell_text(row.cells[1], value, font_size=10)
    return table


def generate_report(analysis: dict, file_name: str) -> bytes:
    today = datetime.now().strftime("%B %d, %Y")
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ── Header ───────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3864')
    pBdr.append(bottom)
    pPr.append(pBdr)
    r1 = p.add_run("GFAM INVESTMENT ANALYSIS")
    r1.bold = True
    r1.font.size = Pt(16)
    r1.font.name = "Arial"
    r1.font.color.rgb = RGBColor.from_string("1F3864")
    r2 = p.add_run(f"    {today}")
    r2.font.size = Pt(10)
    r2.font.name = "Arial"
    r2.font.color.rgb = RGBColor.from_string("888888")

    doc.add_paragraph()

    # ── Recommendation banner ─────────────────────────────────
    rec = clean(analysis.get("recommendation", "CONDITIONAL GO"))
    rec_color = "1F5C2E" if rec == "GO" else "854F0B" if rec == "CONDITIONAL GO" else "A32D2D"
    rec_bg = "EAF3DE" if rec == "GO" else "FAEEDA" if rec == "CONDITIONAL GO" else "FCEBEB"

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), rec_bg)
    pPr.append(shd)
    r = p.add_run(f"  RECOMMENDATION: {rec}  —  {clean(analysis.get('recommendation_rationale', ''))}")
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = "Arial"
    r.font.color.rgb = RGBColor.from_string(rec_color)

    doc.add_paragraph()

    # ── Deal snapshot ─────────────────────────────────────────
    add_section_heading(doc, "Deal Snapshot")
    doc.add_paragraph()
    econ = analysis.get("deal_economics", {})
    add_two_col_table(doc, [
        ("Company", analysis.get("company_name", "Unknown"), "DCE6F1"),
        ("Sector", analysis.get("sector", "Unknown"), "EEF3FB"),
        ("Deal Type", analysis.get("deal_type", "Unknown"), "DCE6F1"),
        ("Asking Price / EV", econ.get("asking_price", "Unknown"), "EEF3FB"),
        ("EV/EBITDA", econ.get("ev_ebitda_multiple", "Unknown"), "DCE6F1"),
        ("Deal Size", econ.get("deal_size", "Unknown"), "EEF3FB"),
        ("Structure", econ.get("structure", "Unknown"), "DCE6F1"),
        ("Capital Product", analysis.get("recommended_capital_product", "Unknown"), "EEF3FB"),
        ("GFAM Fit Score", f"{analysis.get('gfam_fit_score', 'N/A')}/10", "DCE6F1"),
    ])

    doc.add_paragraph()

    # ── Financial snapshot ────────────────────────────────────
    add_section_heading(doc, "Financial Snapshot")
    doc.add_paragraph()
    fin = analysis.get("financial_snapshot", {})
    add_two_col_table(doc, [
        ("Revenue", fin.get("revenue", "Unknown"), "DCE6F1"),
        ("EBITDA", fin.get("ebitda", "Unknown"), "EEF3FB"),
        ("EBITDA Margin", fin.get("ebitda_margin", "Unknown"), "DCE6F1"),
        ("Revenue Growth", fin.get("revenue_growth", "Unknown"), "EEF3FB"),
        ("Net Income", fin.get("net_income", "Unknown"), "DCE6F1"),
        ("Existing Debt", fin.get("debt", "Unknown"), "EEF3FB"),
        ("Recurring Revenue", fin.get("recurring_revenue", "Unknown"), "DCE6F1"),
    ])

    doc.add_paragraph()

    # ── Financial quality ─────────────────────────────────────
    add_section_heading(doc, "Financial Quality Assessment")
    add_body(doc, analysis.get("financial_quality", ""))

    # ── Value creation thesis ─────────────────────────────────
    add_section_heading(doc, "Value Creation Thesis")
    add_body(doc, analysis.get("value_creation_thesis", ""), italic=True, color="1F3864")

    # ── Market context ────────────────────────────────────────
    add_section_heading(doc, "Market Context")
    add_body(doc, analysis.get("market_context", ""))

    # ── GFAM fit ──────────────────────────────────────────────
    add_section_heading(doc, "GFAM Fit")
    add_body(doc, analysis.get("gfam_fit_reason", ""))

    # ── Strengths & Risks side by side ────────────────────────
    add_section_heading(doc, "Key Strengths")
    for s in analysis.get("key_strengths", []):
        add_bullet(doc, s)

    add_section_heading(doc, "Key Risks")
    for r in analysis.get("key_risks", []):
        add_bullet(doc, r)

    red_flags = analysis.get("red_flags", [])
    if red_flags:
        add_section_heading(doc, "Red Flags")
        for f in red_flags:
            add_bullet(doc, f, color="A32D2D")

    # ── Diligence checklist ───────────────────────────────────
    add_section_heading(doc, "Due Diligence Checklist")
    for i, item in enumerate(analysis.get("diligence_checklist", []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(f"{i}. {clean(str(item))}")
        run.font.size = Pt(10)
        run.font.name = "Arial"

    # ── Footer ────────────────────────────────────────────────
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fr = fp.add_run(f"GFAM Investment Analysis  |  Confidential  |  {today}")
    fr.font.size = Pt(8)
    fr.font.name = "Arial"
    fr.font.color.rgb = RGBColor.from_string("999999")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
